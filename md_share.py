#!/usr/bin/env python3

from __future__ import annotations

import argparse
import datetime as dt
import os
import re
import shutil
import subprocess
import sys
import tempfile
import urllib.parse
import urllib.request
from pathlib import Path


ROOT = Path(__file__).resolve().parent
WORD_TEMPLATE = ROOT / "md-to-word-quick-action" / "reference.docx"
HTML_BUILDER = ROOT / "md-to-html-quick-action" / "build_pretty_html.py"
MERMAID_PRERENDER = ROOT / "mermaid-prerender.py"
MARKDOWN_SUFFIXES = {".md", ".markdown"}
HTML_SUFFIXES = {".html", ".htm"}
HTML_THEMES = {"classic", "article", "report", "reading", "interactive"}
LOG_PATH = Path(os.environ.get("MD_SHARE_LOG", Path(tempfile.gettempdir()) / "markdown-share-quick-actions.log"))


def log(message: str) -> None:
    timestamp = dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        with LOG_PATH.open("a", encoding="utf-8") as f:
            f.write(f"[{timestamp}] {message}\n")
    except Exception:
        pass


def find_pandoc() -> str:
    env_value = os.environ.get("PANDOC_BIN")
    candidates = [
        env_value,
        shutil.which("pandoc"),
        "/opt/homebrew/bin/pandoc",
        "/usr/local/bin/pandoc",
        r"C:\Program Files\Pandoc\pandoc.exe",
        r"C:\Users\%USERNAME%\AppData\Local\Pandoc\pandoc.exe",
    ]
    for candidate in candidates:
        if not candidate:
            continue
        expanded = os.path.expandvars(os.path.expanduser(candidate))
        if Path(expanded).exists() or shutil.which(expanded):
            return expanded
    raise RuntimeError("pandoc was not found. Install pandoc first: https://pandoc.org/installing.html")


def find_node() -> str | None:
    candidates = [
        os.environ.get("NODE_BIN"),
        shutil.which("node"),
        "/opt/homebrew/bin/node",
        "/usr/local/bin/node",
        r"C:\Program Files\nodejs\node.exe",
        r"C:\Program Files (x86)\nodejs\node.exe",
    ]
    for candidate in candidates:
        if not candidate:
            continue
        expanded = os.path.expandvars(os.path.expanduser(candidate))
        if Path(expanded).exists() or shutil.which(expanded):
            return expanded
    return None


def find_mermaid_cli() -> str | None:
    npm_prefix = os.environ.get("APPDATA")
    candidates = [
        os.environ.get("MERMAID_CLI"),
        shutil.which("mmdc"),
        "/opt/homebrew/bin/mmdc",
        "/usr/local/bin/mmdc",
        str(Path(npm_prefix) / "npm" / "mmdc.cmd") if npm_prefix else None,
        str(Path(npm_prefix) / "npm" / "mmdc") if npm_prefix else None,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        expanded = os.path.expandvars(os.path.expanduser(candidate))
        if Path(expanded).exists() or shutil.which(expanded):
            return expanded
    return None


def has_mermaid_blocks(path: Path) -> bool:
    try:
        return "```mermaid" in path.read_text(encoding="utf-8", errors="ignore").lower()
    except Exception:
        return False


def check_environment(paths: list[Path], recursive: bool, pandoc_override: str | None = None) -> int:
    ok = True
    print("Markdown Share Quick Actions environment check")
    try:
        pandoc = pandoc_override or find_pandoc()
        print(f"[OK] pandoc: {pandoc}")
    except Exception as exc:
        print(f"[FAIL] pandoc: {exc}")
        ok = False

    files = iter_markdown_files(paths, recursive) if paths else []
    needs_mermaid = any(has_mermaid_blocks(path) for path in files)
    if files:
        print(f"[OK] markdown files: {len(files)}")
    else:
        print("[INFO] markdown files: none supplied; Mermaid checks are dependency-only")

    node = find_node()
    mermaid = find_mermaid_cli()
    if node:
        print(f"[OK] node: {node}")
    elif needs_mermaid:
        print("[FAIL] node: required for Mermaid diagrams. Install Node.js: https://nodejs.org/")
        ok = False
    else:
        print("[INFO] node: not found; only needed for Mermaid diagrams")

    if mermaid:
        print(f"[OK] mermaid cli: {mermaid}")
    elif needs_mermaid:
        print("[FAIL] mermaid cli: required for Mermaid diagrams. Install with: npm install -g @mermaid-js/mermaid-cli")
        ok = False
    else:
        print("[INFO] mermaid cli: not found; only needed for Mermaid diagrams")

    print(f"[INFO] log: {LOG_PATH}")
    return 0 if ok else 1


def preprocess_mermaid(src: Path, tmp_dir: Path, fmt: str) -> Path:
    """Render ```mermaid blocks to images; return path to rewritten MD."""
    if not MERMAID_PRERENDER.exists():
        log(f"mermaid skip: missing prerender script {MERMAID_PRERENDER}")
        return src
    text = src.read_text(encoding="utf-8")
    if "```mermaid" not in text:
        log(f"mermaid skip: no blocks in {src}")
        return src
    out_md = tmp_dir / "mermaid-rendered.md"
    assets_dir = tmp_dir / "mermaid-assets"
    log(f"mermaid start: src={src} fmt={fmt} assets={assets_dir}")
    result = subprocess.run(
        [sys.executable, str(MERMAID_PRERENDER), str(src), str(out_md), str(assets_dir), fmt],
        capture_output=True,
        text=True,
    )
    if result.returncode == 0 and out_md.exists():
        rendered_text = out_md.read_text(encoding="utf-8")
        image_count = rendered_text.count("![图")
        remaining = rendered_text.count("```mermaid")
        log(f"mermaid done: images={image_count} remaining_blocks={remaining}")
        if remaining:
            raise RuntimeError(
                f"Mermaid render incomplete: images={image_count}, remaining_blocks={remaining}. "
                f"Details: {(result.stderr or '').strip()[:800]}"
            )
        return out_md
    log(
        "mermaid failed: "
        f"returncode={result.returncode} stderr={(result.stderr or '').strip()[:800]}"
    )
    raise RuntimeError(f"Mermaid render failed: {(result.stderr or '').strip()[:1200]}")


def postprocess_docx(dest: Path) -> None:
    """Center tables and image captions in a pandoc-generated docx."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn as dqn
        from docx.oxml import OxmlElement
    except ImportError:
        return
    doc = Document(dest)
    for table in doc.tables:
        tblPr = table._tbl.find(dqn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)
        jc = tblPr.find(dqn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(dqn("w:val"), "center")
    for para in doc.paragraphs:
        if para.style.name == "Image Caption":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(dest)


def preprocess_markdown(src: Path, dst: Path) -> None:
    text = src.read_text(encoding="utf-8")
    pattern = re.compile(r"!\[\[([^\]]+)\]\]")

    def replace(match: re.Match[str]) -> str:
        inner = match.group(1).strip()
        if not inner:
            return match.group(0)
        parts = [part.strip() for part in inner.split("|") if part.strip()]
        target = parts[0]
        alt = ""
        if len(parts) >= 2 and not parts[1].isdigit():
            alt = parts[1]
        if not alt:
            alt = Path(target).stem
        return f"![{alt}]({target})"

    dst.write_text(pattern.sub(replace, text), encoding="utf-8")


def output_path(src: Path, suffix: str, output_dir: Path | None) -> Path:
    base_dir = output_dir if output_dir else src.parent
    base_dir.mkdir(parents=True, exist_ok=True)
    return base_dir / f"{src.stem}.{suffix}"


def themed_output_path(src: Path, theme: str, output_dir: Path | None) -> Path:
    if theme == "classic":
        return output_path(src, "html", output_dir)
    base_dir = output_dir if output_dir else src.parent
    base_dir.mkdir(parents=True, exist_ok=True)
    return base_dir / f"{src.stem}.{theme}.html"


def convert_word(src: Path, output_dir: Path | None, pandoc: str) -> Path:
    if not WORD_TEMPLATE.exists():
        raise RuntimeError(f"Missing Word reference template: {WORD_TEMPLATE}")
    dest = output_path(src, "docx", output_dir)
    log(f"word start: src={src} dest={dest} root={ROOT}")
    with tempfile.TemporaryDirectory(prefix="md-share-word.") as tmp:
        tmp_dir = Path(tmp)
        prepared = tmp_dir / "prepared.md"
        preprocess_markdown(src, prepared)
        ready = preprocess_mermaid(prepared, tmp_dir, "png")
        cmd = [
            pandoc,
            str(ready),
            "-o",
            str(dest),
            f"--reference-doc={WORD_TEMPLATE}",
            f"--resource-path={src.parent}",
        ]
        subprocess.run(cmd, check=True)
    postprocess_docx(dest)
    log(f"word done: dest={dest} exists={dest.exists()} size={dest.stat().st_size if dest.exists() else 0}")
    return dest


def convert_html(src: Path, output_dir: Path | None, pandoc: str, theme: str = "classic") -> Path:
    if not HTML_BUILDER.exists():
        raise RuntimeError(f"Missing HTML builder: {HTML_BUILDER}")
    if theme not in HTML_THEMES:
        raise RuntimeError(f"Unknown HTML theme: {theme}. Choose one of: {', '.join(sorted(HTML_THEMES))}")
    dest = themed_output_path(src, theme, output_dir)
    log(f"html start: src={src} dest={dest} theme={theme} root={ROOT}")
    with tempfile.TemporaryDirectory(prefix="md-share-html.") as tmp:
        tmp_dir = Path(tmp)
        prepared = tmp_dir / "prepared.md"
        rendered = tmp_dir / "rendered.html"
        preprocess_markdown(src, prepared)
        ready = preprocess_mermaid(prepared, tmp_dir, "svg")
        cmd = [
            pandoc,
            str(ready),
            "-o",
            str(rendered),
            "--standalone",
            "--embed-resources",
            f"--resource-path={src.parent}",
            "--metadata",
            f"title={src.stem}",
        ]
        if theme == "interactive":
            cmd.extend(["--toc", "--toc-depth=3"])
        subprocess.run(cmd, check=True)
        subprocess.run([sys.executable, str(HTML_BUILDER), str(rendered), str(dest), src.stem, theme], check=True)
    log(f"html done: dest={dest} exists={dest.exists()} size={dest.stat().st_size if dest.exists() else 0}")
    return dest


def convert_any_to_md(src: Path, output_dir: Path | None) -> Path:
    dest = output_path(src, "md", output_dir)
    log(f"to-md start: src={src} dest={dest}")

    # FIX: Try pandoc first for format-preserving conversion (headings, tables, images)
    ext = src.suffix.lower()
    pandoc_supported = {".docx", ".xlsx", ".pptx", ".pdf", ".html", ".htm", ".odt", ".rtf", ".epub"}
    if ext in pandoc_supported:
        try:
            pandoc = find_pandoc()
            subprocess.run(
                [pandoc, str(src), "-t", "markdown", "-o", str(dest), "--wrap=none"],
                check=True,
                capture_output=True,
                text=True,
            )
            log(f"to-md done via pandoc: {dest}")
            return dest
        except Exception as exc:
            log(f"pandoc failed for {src}: {exc}, falling back to markitdown")

    # Fallback to markitdown (text-only extraction)
    try:
        from markitdown import MarkItDown
    except ImportError as exc:
        raise RuntimeError("markitdown is not installed. Install with: python3 -m pip install 'markitdown[all]'") from exc
    converter = MarkItDown()
    result = converter.convert(str(src))
    content = result.text_content or ""
    if not content.strip():
        raise RuntimeError(f"Conversion produced empty Markdown: {src}")
    dest.write_text(content, encoding="utf-8")
    log(f"to-md done via markitdown: {dest} exists={dest.exists()} size={dest.stat().st_size if dest.exists() else 0}")
    return dest


def convert_html_to_md(src: Path, output_dir: Path | None) -> Path:
    dest = output_path(src, "md", output_dir)
    log(f"html-to-md start: src={src} dest={dest}")
    html_text = src.read_text(encoding="utf-8", errors="replace")
    md_text = html_to_markdown(html_text)
    if not md_text.strip():
        raise RuntimeError(f"Conversion produced empty Markdown: {src}")
    dest.write_text(md_text, encoding="utf-8")
    log(f"html-to-md done: dest={dest} exists={dest.exists()} size={dest.stat().st_size if dest.exists() else 0}")
    return dest


def convert_url_to_md(url: str, output_dir: Path | None, mode: str) -> Path:
    name = url_to_stem(url)
    base_dir = output_dir if output_dir else Path.cwd()
    base_dir.mkdir(parents=True, exist_ok=True)
    dest = base_dir / f"{name}.md"
    log(f"url-to-md start: url={url} dest={dest} mode={mode}")
    html_text = fetch_url(url)
    if mode in {"article", "auto"}:
        html_text = extract_article_html(html_text, url) or html_text
    md_text = html_to_markdown(html_text)
    if not md_text.strip():
        raise RuntimeError(f"Conversion produced empty Markdown: {url}")
    dest.write_text(md_text, encoding="utf-8")
    log(f"url-to-md done: dest={dest} exists={dest.exists()} size={dest.stat().st_size if dest.exists() else 0}")
    return dest


def fetch_url(url: str) -> str:
    req = urllib.request.Request(url, headers={"User-Agent": "MarkdownShareQuickActions/1.0"})
    with urllib.request.urlopen(req, timeout=30) as resp:
        raw = resp.read()
        charset = resp.headers.get_content_charset() or "utf-8"
        return raw.decode(charset, errors="replace")


def extract_article_html(html_text: str, url: str) -> str | None:
    try:
        import trafilatura
    except ImportError:
        return None
    return trafilatura.extract(
        html_text,
        url=url,
        output_format="html",
        include_comments=False,
        include_tables=True,
        favor_recall=False,
    )


def html_to_markdown(html_text: str) -> str:
    try:
        from html_to_markdown import convert
        result = convert(html_text)
        return getattr(result, "content", result if isinstance(result, str) else "")
    except ImportError:
        pass
    try:
        from markdownify import markdownify
    except ImportError as exc:
        raise RuntimeError("Install html-to-markdown or markdownify for HTML to Markdown conversion.") from exc
    return markdownify(
        html_text,
        heading_style="ATX",
        bullets="-",
        strip=["script", "style", "nav", "footer", "aside", "iframe", "form"],
    )


def url_to_stem(url: str) -> str:
    parsed = urllib.parse.urlparse(url)
    name = Path(parsed.path.rstrip("/") or parsed.netloc).stem or "fetched"
    return re.sub(r"[^A-Za-z0-9._-]+", "-", name).strip("-") or "fetched"


def iter_markdown_files(paths: list[Path], recursive: bool) -> list[Path]:
    files: list[Path] = []
    for path in paths:
        expanded = path.expanduser().resolve()
        if expanded.is_dir() and recursive:
            files.extend(p for p in expanded.rglob("*") if p.is_file() and p.suffix.lower() in MARKDOWN_SUFFIXES)
        elif expanded.is_file() and expanded.suffix.lower() in MARKDOWN_SUFFIXES:
            files.append(expanded)
    return sorted(dict.fromkeys(files), key=lambda p: str(p).lower())


def iter_files(paths: list[Path], recursive: bool, suffixes: set[str] | None = None) -> list[Path]:
    files: list[Path] = []
    for path in paths:
        expanded = path.expanduser().resolve()
        if expanded.is_dir() and recursive:
            candidates = [p for p in expanded.rglob("*") if p.is_file()]
        elif expanded.is_file():
            candidates = [expanded]
        else:
            candidates = []
        for candidate in candidates:
            if suffixes is None or candidate.suffix.lower() in suffixes:
                files.append(candidate)
    return sorted(dict.fromkeys(files), key=lambda p: str(p).lower())


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert Markdown, HTML, and common document files.")
    parser.add_argument("format", choices=["word", "html", "both", "to-md", "html-to-md", "url-to-md"], help="Output format.")
    parser.add_argument("paths", nargs="*", help="Files, URLs, or directories when --recursive is used.")
    parser.add_argument("-o", "--output-dir", type=Path, default=None, help="Optional output directory.")
    parser.add_argument("--recursive", action="store_true", help="Scan directories recursively for Markdown files.")
    parser.add_argument("--pandoc", default=None, help="Path to pandoc. Overrides auto detection.")
    parser.add_argument("--theme", choices=sorted(HTML_THEMES), default="classic", help="HTML theme. classic preserves the original right-click output.")
    parser.add_argument("--url-mode", choices=["auto", "article", "structured"], default="article", help="URL to Markdown mode.")
    parser.add_argument("--check", action="store_true", help="Check dependencies and selected files without converting.")
    args = parser.parse_args()

    try:
        log(f"invoke: argv={sys.argv} cwd={Path.cwd()}")
        path_args = [Path(path) for path in args.paths]
        if args.check:
            return check_environment(path_args, args.recursive, args.pandoc)
        if not args.paths:
            print("No input supplied.", file=sys.stderr)
            return 2

        outputs: list[Path] = []
        output_dir = args.output_dir.expanduser().resolve() if args.output_dir else None
        if args.format == "url-to-md":
            for item in args.paths:
                outputs.append(convert_url_to_md(item, output_dir, args.url_mode))
        elif args.format == "to-md":
            files = iter_files(path_args, args.recursive, None)
            if not files:
                print("No files found.", file=sys.stderr)
                return 2
            for src in files:
                if src.suffix.lower() in MARKDOWN_SUFFIXES:
                    continue
                if src.suffix.lower() in HTML_SUFFIXES:
                    outputs.append(convert_html_to_md(src, output_dir))
                else:
                    outputs.append(convert_any_to_md(src, output_dir))
        elif args.format == "html-to-md":
            files = iter_files(path_args, args.recursive, HTML_SUFFIXES)
            if not files:
                print("No HTML files found.", file=sys.stderr)
                return 2
            for src in files:
                outputs.append(convert_html_to_md(src, output_dir))
        else:
            pandoc = args.pandoc or find_pandoc()
            files = iter_markdown_files(path_args, args.recursive)
            if not files:
                print("No Markdown files found.", file=sys.stderr)
                return 2
            for src in files:
                if args.format in {"word", "both"}:
                    outputs.append(convert_word(src, output_dir, pandoc))
                if args.format in {"html", "both"}:
                    outputs.append(convert_html(src, output_dir, pandoc, args.theme))

        for path in outputs:
            print(path)
        return 0
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
